"""
DOCX file generator for the Turbo Test File Generator project.
"""

from docx import Document

from generators.base_generator import BaseGenerator
from metadata.builder import build_metadata
from metadata.sidecar_writer import write_sidecar_metadata
from utils.helpers import build_file_path, ensure_directory
from utils.random_data import (
    get_random_paragraph,
    get_random_tags,
    get_random_title,
)


def apply_docx_core_properties(document, metadata):
    """
    Apply supported metadata fields to DOCX core properties.

    Args:
        document: python-docx Document instance.
        metadata: Metadata dictionary.
    """
    core_props = document.core_properties

    core_props.author = str(metadata.get("author", ""))
    core_props.category = str(metadata.get("category", ""))
    core_props.comments = str(metadata.get("comments", ""))
    core_props.identifier = str(metadata.get("test_id", ""))
    core_props.keywords = ", ".join(metadata.get("tags", []))
    core_props.last_modified_by = str(metadata.get("author", ""))
    core_props.subject = str(metadata.get("file_type", ""))
    core_props.title = str(metadata.get("file_name", ""))


class DocxGenerator(BaseGenerator):
    """
    Generator for DOCX files.
    """

    def __init__(
        self,
        output_dir,
        use_sidecar_metadata=True,
        use_embedded_metadata=True,
    ):
        """
        Initialize the DOCX generator with shared settings.
        """
        super().__init__(
            output_dir=output_dir,
            file_type="docx",
            use_sidecar_metadata=use_sidecar_metadata,
            use_embedded_metadata=use_embedded_metadata,
        )

    def generate(self, filename, metadata=None, **kwargs):
        """
        Generate a DOCX file and return information about it.

        Args:
            filename: Name of the DOCX file to generate.
            metadata: Optional custom metadata dictionary.
            **kwargs: Extra options such as output_dir or section_count.

        Returns:
            A dictionary describing the generated file.
        """
        output_dir = kwargs.get("output_dir", self.output_dir)
        section_count = kwargs.get("section_count", 3)

        if not isinstance(section_count, int) or isinstance(section_count, bool) or section_count < 1:
            raise ValueError("section_count must be a positive integer.")

        ensure_directory(output_dir)

        file_path = build_file_path(output_dir, filename)
        built_metadata = build_metadata(filename, self.file_type, metadata)

        document = Document()
        apply_docx_core_properties(document, built_metadata)

        title = get_random_title()
        document.add_heading(title, level=0)

        tags = ", ".join(get_random_tags())
        document.add_paragraph(f"Tags: {tags}")

        for _ in range(section_count):
            section_title = get_random_title()
            document.add_heading(section_title, level=1)

            document.add_paragraph(get_random_paragraph())
            document.add_paragraph(get_random_paragraph())

            bullet_items = get_random_tags(count=3)
            for item in bullet_items:
                document.add_paragraph(item, style="List Bullet")

        document.save(file_path)

        sidecar_path = None

        if self.use_sidecar_metadata:
            sidecar_path = write_sidecar_metadata(file_path, built_metadata)

        return {
            "file_name": filename,
            "file_type": self.file_type,
            "file_path": str(file_path),
            "metadata": built_metadata,
            "size_bytes": file_path.stat().st_size,
            "section_count": section_count,
            "sidecar_metadata_enabled": self.use_sidecar_metadata,
            "sidecar_path": str(sidecar_path) if sidecar_path else None,
            "embedded_metadata_enabled": self.use_embedded_metadata,
        }
